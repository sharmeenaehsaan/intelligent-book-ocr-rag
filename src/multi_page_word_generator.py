from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from src.word_generator import (
    configure_document,
    merge_paragraph_lines,
    group_code_lines,
    clean_ocr_text,
    add_code_block,
    set_run_font
)


def add_visual_element(
    document,
    element,
    page_width
):
    """
    Inserts either a table crop or a figure crop.

    element must contain:
        type: "table" or "figure"
        path: image path
        bbox: (x, y, width, height)
    """

    image_path = element["path"]
    x, y, width, height = element["bbox"]

    paragraph = document.add_paragraph()

    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(4)

    run = paragraph.add_run()

    # Estimate the picture width based on its width
    # relative to the original photographed page.
    width_ratio = width / max(page_width, 1)

    if element["type"] == "table":
        # Tables usually need more width to remain readable.
        width_inches = max(
            3.0,
            min(7.0, width_ratio * 7.0)
        )
    else:
        # Figures and screenshots can usually be smaller.
        width_inches = max(
            1.8,
            min(5.8, width_ratio * 6.5)
        )

    run.add_picture(
        image_path,
        width=Inches(width_inches)
    )


def prepare_visual_elements(
    saved_tables,
    saved_figures
):
    """
    Combines tables and figures into one reading-order list.
    """

    visual_elements = []

    for table in saved_tables:
        visual_elements.append({
            "type": "table",
            "path": table["path"],
            "bbox": table["bbox"]
        })

    for figure in saved_figures:
        visual_elements.append({
            "type": "figure",
            "path": figure["path"],
            "bbox": figure["bbox"]
        })

    visual_elements.sort(
        key=lambda element: (
            element["bbox"][1],
            element["bbox"][0]
        )
    )

    return visual_elements


def add_heading(
    document,
    text,
    level=1
):
    paragraph = document.add_heading(
        text,
        level=level
    )

    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.space_before = Pt(7)
    paragraph.paragraph_format.space_after = Pt(3)

    return paragraph


def add_normal_paragraph(
    document,
    text
):
    paragraph = document.add_paragraph()

    paragraph.paragraph_format.first_line_indent = (
        Inches(0.08)
    )

    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.line_spacing = 1.05

    run = paragraph.add_run(text)

    set_run_font(
        run,
        "Times New Roman",
        10.5
    )

    return paragraph


def insert_visuals_before_position(
    document,
    visual_elements,
    inserted_visuals,
    current_top,
    page_width
):
    """
    Inserts all tables or figures located above the current
    OCR block.
    """

    for visual_index, element in enumerate(
        visual_elements
    ):
        if visual_index in inserted_visuals:
            continue

        element_top = element["bbox"][1]

        if element_top <= current_top:
            add_visual_element(
                document,
                element,
                page_width
            )

            inserted_visuals.add(
                visual_index
            )


def add_page_content(
    document,
    page_data
):
    classified_blocks = page_data.get(
        "classified_blocks",
        []
    )

    saved_tables = page_data.get(
        "saved_tables",
        []
    )

    saved_figures = page_data.get(
        "saved_figures",
        []
    )

    page_width = page_data.get(
        "page_width",
        1000
    )

    blocks = merge_paragraph_lines(
        classified_blocks
    )

    blocks = group_code_lines(
        blocks
    )

    blocks = sorted(
        blocks,
        key=lambda block: (
            block.get("top", 0),
            block.get("left", 0)
        )
    )

    visual_elements = prepare_visual_elements(
        saved_tables,
        saved_figures
    )

    inserted_visuals = set()

    skipped_types = {
    "page_number",
    "inside_table",
    "inside_figure"
}

    for block in blocks:
        block_top = block.get(
            "top",
            0
        )

        insert_visuals_before_position(
            document=document,
            visual_elements=visual_elements,
            inserted_visuals=inserted_visuals,
            current_top=block_top,
            page_width=page_width
        )

        block_type = block.get(
            "type",
            "paragraph"
        )

        if block_type in skipped_types:
            continue

        text = clean_ocr_text(
            block.get("text", "")
        )

        if not text:
            continue

        if block_type == "heading":
            add_heading(
                document,
                text,
                level=1
            )

        elif block_type == "subheading":
            add_heading(
                document,
                text,
                level=2
            )

        elif block_type == "code":
            add_code_block(
                document,
                block,
                page_width=page_width
            )

        else:
            add_normal_paragraph(
                document,
                text
            )

    # Insert visual items located below the final OCR block.
    for visual_index, element in enumerate(
        visual_elements
    ):
        if visual_index in inserted_visuals:
            continue

        add_visual_element(
            document,
            element,
            page_width
        )

        inserted_visuals.add(
            visual_index
        )


def build_multi_page_document(
    processed_pages,
    output_path="output/complete_book.docx",
    preserve_input_pages=True
):
    """
    Builds one Word document from all processed book pages.

    preserve_input_pages=True:
        adds a page break after every photographed page.

    preserve_input_pages=False:
        allows content to flow continuously in Word.
    """

    if not processed_pages:
        raise ValueError(
            "No processed pages were supplied."
        )

    output_path = Path(
        output_path
    )

    output_path.parent.mkdir(
        parents=True,
        exist_ok=True
    )

    document = Document()

    configure_document(
        document
    )

    total_pages = len(
        processed_pages
    )

    for page_index, page_data in enumerate(
        processed_pages,
        start=1
    ):
        source_image = page_data.get(
            "source_image",
            "Unknown image"
        )

        print(
            f"Writing page "
            f"{page_index}/{total_pages}: "
            f"{Path(source_image).name}"
        )

        add_page_content(
            document,
            page_data
        )

        if (
            preserve_input_pages
            and page_index < total_pages
        ):
            document.add_page_break()

    document.save(
        output_path
    )

    print(
        f"\nComplete Word document saved at:\n"
        f"{output_path}"
    )