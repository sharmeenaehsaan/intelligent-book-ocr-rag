import re
from pathlib import Path
from statistics import median
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


def clean_basic_text(text):
    text = text.strip()

    replacements = {
        " .": ".",
        " ,": ",",
        " ;": ";",
        " :": ":",
        "( ": "(",
        " )": ")"
    }

    for old, new in replacements.items():
        text = text.replace(old, new)

    return text


def fix_spacing(text):
    """
    Fixes common OCR spacing problems such as:

    application.However -> application. However
    thread,not          -> thread, not
    applet,you          -> applet, you
    method.The          -> method. The
    """

    text = re.sub(
        r"([a-z0-9])\.([A-Z])",
        r"\1. \2",
        text
    )

    text = re.sub(
        r"([a-zA-Z0-9]),([a-zA-Z])",
        r"\1, \2",
        text
    )

    text = re.sub(
        r"([a-zA-Z0-9]);([a-zA-Z])",
        r"\1; \2",
        text
    )

    text = re.sub(
        r"([a-zA-Z0-9]):([a-zA-Z])",
        r"\1: \2",
        text
    )

    text = re.sub(
        r"\)\.([A-Z])",
        r"). \1",
        text
    )

    text = re.sub(
        r"\)([A-Za-z])",
        r") \1",
        text
    )

    text = re.sub(
        r"([.!?])([A-Z])",
        r"\1 \2",
        text
    )

    text = re.sub(
        r"\s+",
        " ",
        text
    )

    return text.strip()


def clean_ocr_text(text):
    text = clean_basic_text(text)
    text = fix_spacing(text)
    return text


def should_join_lines(current_block, next_block):
    if current_block["type"] != "paragraph":
        return False

    if next_block["type"] != "paragraph":
        return False

    current_text = current_block["text"].strip()
    next_text = next_block["text"].strip()

    if not current_text or not next_text:
        return False

    vertical_gap = (
        next_block["top"]
        - current_block["bottom"]
    )

    typical_height = max(
        current_block["height"],
        next_block["height"]
    )

    if vertical_gap > typical_height * 0.8:
        return False

    indentation_difference = abs(
        next_block["left"]
        - current_block["left"]
    )

    if indentation_difference > 45:
        return False

    return True


def merge_paragraph_lines(classified_blocks):
    blocks = sorted(
        classified_blocks,
        key=lambda block: (
            block["top"],
            block["left"]
        )
    )

    merged_blocks = []
    index = 0

    while index < len(blocks):
        block = blocks[index]

        if block["type"] != "paragraph":
            copied_block = block.copy()
            copied_block["text"] = clean_ocr_text(
                copied_block.get("text", "")
            )

            merged_blocks.append(
                copied_block
            )

            index += 1
            continue

        paragraph_text = clean_ocr_text(
            block["text"]
        )

        paragraph_block = block.copy()

        next_index = index + 1

        while next_index < len(blocks):
            previous_line = blocks[next_index - 1]
            next_block = blocks[next_index]

            if not should_join_lines(
                previous_line,
                next_block
            ):
                break

            next_text = clean_ocr_text(
                next_block["text"]
            )

            if paragraph_text.endswith("-"):
                paragraph_text = (
                    paragraph_text[:-1]
                    + next_text
                )
            else:
                paragraph_text += " " + next_text

            paragraph_block["bottom"] = (
                next_block["bottom"]
            )

            paragraph_block["height"] = (
                paragraph_block["bottom"]
                - paragraph_block["top"]
            )

            next_index += 1

        paragraph_block["text"] = fix_spacing(
            paragraph_text
        )

        merged_blocks.append(
            paragraph_block
        )

        index = next_index

    return merged_blocks


def group_code_lines(blocks):
    grouped_blocks = []
    index = 0

    while index < len(blocks):
        block = blocks[index]

        if block.get("type") != "code":
            grouped_blocks.append(block)
            index += 1
            continue

        code_lines = [{
            "text": clean_basic_text(
                block.get("text", "")
            ),
            "left": block.get("left", 0),
            "right": block.get("right", 0),
            "top": block.get("top", 0),
            "bottom": block.get("bottom", 0),
            "width": block.get("width", 0),
            "height": block.get("height", 0)
        }]

        grouped_block = block.copy()
        next_index = index + 1

        while next_index < len(blocks):
            next_block = blocks[next_index]

            if next_block.get("type") != "code":
                break

            previous_block = blocks[
                next_index - 1
            ]

            vertical_gap = (
                next_block.get("top", 0)
                - previous_block.get("bottom", 0)
            )

            typical_height = max(
                previous_block.get("height", 1),
                next_block.get("height", 1)
            )

            if (
                vertical_gap
                > typical_height * 1.4
            ):
                break

            code_lines.append({
                "text": clean_basic_text(
                    next_block.get("text", "")
                ),
                "left": next_block.get("left", 0),
                "right": next_block.get("right", 0),
                "top": next_block.get("top", 0),
                "bottom": next_block.get(
                    "bottom",
                    0
                ),
                "width": next_block.get("width", 0),
                "height": next_block.get(
                    "height",
                    0
                )
            })

            next_index += 1

        grouped_block["code_lines"] = (
            code_lines
        )

        grouped_block["text"] = "\n".join(
            line["text"]
            for line in code_lines
        )

        grouped_block["left"] = min(
            line["left"]
            for line in code_lines
        )

        grouped_block["right"] = max(
            line["right"]
            for line in code_lines
        )

        grouped_block["top"] = min(
            line["top"]
            for line in code_lines
        )

        grouped_block["bottom"] = max(
            line["bottom"]
            for line in code_lines
        )

        grouped_block["width"] = (
            grouped_block["right"]
            - grouped_block["left"]
        )

        grouped_block["height"] = (
            grouped_block["bottom"]
            - grouped_block["top"]
        )

        grouped_blocks.append(
            grouped_block
        )

        index = next_index

    return grouped_blocks

def set_cell_shading(cell, fill):
    cell_properties = (
        cell._tc.get_or_add_tcPr()
    )

    shading = OxmlElement("w:shd")

    shading.set(
        qn("w:fill"),
        fill
    )

    cell_properties.append(shading)


def set_run_font(run, font_name, font_size):
    run.font.name = font_name
    run.font.size = Pt(font_size)

    run_properties = run._element.get_or_add_rPr()

    fonts = run_properties.rFonts

    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        run_properties.insert(0, fonts)

    fonts.set(
        qn("w:ascii"),
        font_name
    )

    fonts.set(
        qn("w:hAnsi"),
        font_name
    )


def add_code_block(
    document,
    code_block,
    page_width=1000
):
    paragraph = document.add_paragraph()

    paragraph.paragraph_format.left_indent = (
        Inches(0.25)
    )

    paragraph.paragraph_format.right_indent = (
        Inches(0.10)
    )

    paragraph.paragraph_format.first_line_indent = (
        Inches(0)
    )

    paragraph.paragraph_format.space_before = (
        Pt(2)
    )

    paragraph.paragraph_format.space_after = (
        Pt(4)
    )

    paragraph.paragraph_format.line_spacing = 1

    if isinstance(code_block, str):
        code_lines = [
            {
                "text": line,
                "left": 0,
                "width": 0
            }
            for line in code_block.splitlines()
        ]
    else:
        code_lines = code_block.get(
            "code_lines",
            []
        )

        if not code_lines:
            code_lines = [
                {
                    "text": line,
                    "left": code_block.get(
                        "left",
                        0
                    ),
                    "width": code_block.get(
                        "width",
                        0
                    )
                }
                for line in code_block.get(
                    "text",
                    ""
                ).splitlines()
            ]

    if not code_lines:
        return paragraph

    base_left = min(
        line.get("left", 0)
        for line in code_lines
    )

    character_widths = []

    for line in code_lines:
        line_text = line.get(
            "text",
            ""
        ).strip()

        line_width = line.get(
            "width",
            0
        )

        if (
            len(line_text) >= 4
            and line_width > 0
        ):
            character_widths.append(
                line_width / len(line_text)
            )

    if character_widths:
        typical_character_width = median(
            character_widths
        )
    else:
        typical_character_width = max(
            page_width * 0.007,
            6
        )

    for line_index, line in enumerate(
        code_lines
    ):
        if line_index > 0:
            paragraph.add_run().add_break()

        line_text = line.get(
            "text",
            ""
        )

        relative_left = max(
            0,
            line.get("left", base_left)
            - base_left
        )

        raw_indent_spaces = (
            relative_left
            / max(typical_character_width, 1)
        )

        if raw_indent_spaces < 0.75:
            indent_spaces = 0
        else:
            indent_spaces = int(
                round(
                    raw_indent_spaces / 2
                ) * 2
            )

        indent_spaces = min(
            indent_spaces,
            40
        )

        formatted_line = (
            " " * indent_spaces
            + line_text
        )

        run = paragraph.add_run(
            formatted_line
        )

        set_run_font(
            run,
            "Courier New",
            9
        )

    return paragraph


def add_figure(
    document,
    figure_path,
    width_inches=2.3
):
    paragraph = document.add_paragraph()

    paragraph.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )

    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(4)

    run = paragraph.add_run()

    run.add_picture(
        figure_path,
        width=Inches(width_inches)
    )


def configure_document(document):
    section = document.sections[0]

    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)

    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    normal_style = document.styles["Normal"]

    normal_style.font.name = (
        "Times New Roman"
    )

    normal_style.font.size = Pt(10.5)

    normal_format = (
        normal_style.paragraph_format
    )

    normal_format.space_before = Pt(0)
    normal_format.space_after = Pt(2)
    normal_format.line_spacing = 1.05

    heading_style = (
        document.styles["Heading 1"]
    )

    heading_style.font.name = (
        "Times New Roman"
    )

    heading_style.font.size = Pt(14)
    heading_style.font.bold = True

    heading_format = (
        heading_style.paragraph_format
    )

    heading_format.space_before = Pt(8)
    heading_format.space_after = Pt(3)


def build_word_document(
    classified_blocks,
    saved_figures,
    output_path="output/book.docx"
):
    output_path = Path(output_path)

    output_path.parent.mkdir(
        parents=True,
        exist_ok=True
    )

    document = Document()

    configure_document(document)

    blocks = merge_paragraph_lines(
        classified_blocks
    )

    blocks = group_code_lines(
        blocks
    )

    figure_data = sorted(
        saved_figures,
        key=lambda figure: (
            figure["bbox"][1],
            figure["bbox"][0]
        )
    )

    inserted_figures = set()

    for block in blocks:
        block_type = block["type"]

        text = clean_ocr_text(
            block.get("text", "")
        )

        for figure_index, figure in enumerate(
            figure_data
        ):
            if figure_index in inserted_figures:
                continue

            figure_top = figure["bbox"][1]

            if figure_top <= block["top"]:
                add_figure(
                    document,
                    figure["path"],
                    width_inches=2.3
                )

                inserted_figures.add(
                    figure_index
                )

        if block_type in {
            "header",
            "page_number",
            "inside_figure",
            "figure_text"
        }:
            continue

        if not text:
            continue

        if block_type == "heading":
            paragraph = document.add_heading(
                text,
                level=1
            )

            paragraph.paragraph_format.keep_with_next = (
                True
            )

        elif block_type == "paragraph":
            paragraph = document.add_paragraph()

            paragraph.paragraph_format.first_line_indent = (
                Inches(0.08)
            )

            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(2)
            paragraph.paragraph_format.line_spacing = 1.05

            run = paragraph.add_run(text)

            set_run_font(
                run,
                "Times New Roman",
                10.5
            )

        elif block_type == "code":
            add_code_block(
                document,
                block
            )

    for figure_index, figure in enumerate(
        figure_data
    ):
        if figure_index not in inserted_figures:
            add_figure(
                document,
                figure["path"],
                width_inches=2.3
            )

    document.save(
        output_path
    )

    print(
        f"\nSaved Word document: "
        f"{output_path}"
    )